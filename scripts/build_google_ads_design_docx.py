#!/usr/bin/env python3
"""Build docs/engineering/google-ads-api-tool-design-document.docx from source content.

Requires: pip install python-docx

Optional: LibreOffice (`soffice`) on PATH to also emit a legacy .doc file.
"""

from __future__ import annotations

import platform
import subprocess
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "docs" / "engineering" / "google-ads-api-tool-design-document.docx"
DOC_PATH = ROOT / "docs" / "engineering" / "google-ads-api-tool-design-document.doc"
IMAGE_PATH = ROOT / "docs" / "engineering" / "assets" / "google-ads-report-prototype.png"
VENDOR = ROOT / ".vendor" / "python-docx"

if VENDOR.is_dir():
    sys.path.insert(0, str(VENDOR))

try:
    from docx import Document
    from docx.shared import Inches, Pt
except ImportError as exc:  # pragma: no cover
    raise SystemExit(
        "Missing python-docx. Install with: pip install python-docx\n"
        f"Or: pip install -t {VENDOR} python-docx"
    ) from exc


def add_bullets(doc: Document, items: list[str]) -> None:
    for line in items:
        p = doc.add_paragraph(line, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(11)


def add_table(doc: Document, headers: tuple[str, ...], rows: list[tuple[str, ...]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text


def main() -> None:
    doc = Document()

    t = doc.add_heading("Google Ads API — Tool design document", 0)
    for r in t.runs:
        r.font.size = Pt(22)

    meta = doc.add_paragraph()
    meta.add_run("Product: ").bold = True
    meta.add_run("Duct (https://getduct.ai)\n")
    meta.add_run("Document version: ").bold = True
    meta.add_run("1.0\n")
    meta.add_run("Last updated: ").bold = True
    meta.add_run("April 2026\n\n")
    meta.add_run(
        "This document follows the structure Google’s token application materials use as a reference. "
        "Submit as .doc or .docx per Google’s form. If your tool is externally accessible, include "
        "screenshots or mock-ups (see §7)."
    )

    doc.add_heading("1. Company name", level=1)
    p = doc.add_paragraph()
    p.add_run("Duct").bold = True
    p.add_run(" (public site: https://getduct.ai)")
    doc.add_paragraph(
        "If your legal entity differs (e.g. “Alleviate Lab LLC”), add it here in parentheses.",
        style="Intense Quote",
    )

    doc.add_heading("2. Business model", level=1)
    doc.add_paragraph(
        "Duct is a software product that helps advertisers understand Google Ads performance through "
        "structured reports and goal-oriented analysis (for example: efficiency, ROAS, scaling, or spend audits). "
        "The product is developed and operated by our company; we are not a traditional agency managing "
        "third-party ad accounts as the primary business, though end users (including marketers and agencies) "
        "may use Duct with their own Google Ads accounts where permitted."
    )
    doc.add_paragraph(
        "We use the Google Ads API only to read performance and account metadata that the authenticated user "
        "is authorized to access. We do not use the API to create Google Ads accounts, create or edit campaigns "
        "or ads, or to run Keyword Planner as a service."
    )

    doc.add_heading("3. Tool access / use", level=1)
    doc.add_paragraph("Who uses the tool", style="Heading 3")
    add_bullets(
        doc,
        [
            "Primary: Authenticated users (internal team during development; external customers as the product "
            "matures) who connect their Google Ads account via OAuth and run reports on demand.",
            "Access model: Users initiate a report generation flow in the web app: they choose data sources "
            "(Google Ads), select account (when multiple are accessible), analysis goal, optional business "
            "context, and date range. The backend calls the Google Ads API with that user’s refresh token for "
            "that session/request and returns a JSON brief rendered in the UI. Users may save reports for local "
            "viewing in the app (e.g. browser storage for demos); production deployments may persist reports "
            "according to our data policy.",
        ],
    )
    doc.add_paragraph("What we do not do (today)", style="Heading 3")
    add_bullets(
        doc,
        [
            "We do not expose our Google Ads developer token to arbitrary third-party tools.",
            "We do not run a scheduled batch job (e.g. hourly) that mutates Google Ads entities for inventory or "
            "stock status. Duct’s current implementation is read-only toward Google Ads.",
            "We do not grant agency partners direct login to our app unless they are normal product users; any "
            "sharing of exported insights (PDF/screenshot/email) is outside the API tool’s access boundary.",
        ],
    )
    doc.add_paragraph(
        "Adjust the “internal vs external” sentences above to match exactly what you selected on the application.",
        style="Intense Quote",
    )

    doc.add_heading("4. Tool design", level=1)
    doc.add_paragraph("Architecture (high level)", style="Heading 3")
    add_bullets(
        doc,
        [
            "Web client (Next.js) — Pages for connections (OAuth), generate report (wizard: sources → configuration "
            "→ review → generate), and report viewing.",
            "Application backend (FastAPI) — Serves REST APIs under /api/…, validates API key on protected routes, "
            "performs OAuth redirects for Google (browser flow), and orchestrates fetch → brief → optional LLM synthesis.",
            "Google Ads API — Read-only access using GoogleAdsClient with developer token, OAuth client id/secret, "
            "user refresh token, and optional login_customer_id (MCC) when listing or querying client accounts under a manager.",
        ],
    )
    doc.add_paragraph("Data flow", style="Heading 3")
    add_bullets(
        doc,
        [
            "User completes Google OAuth (scope: https://www.googleapis.com/auth/adwords). Refresh token is stored "
            "in the browser for the demo-style flow (sessionStorage) and sent to our backend only when requesting "
            "account list or generation; production should move to server-side token storage tied to user identity.",
            "Account listing: Backend calls CustomerService.list_accessible_customers, then GoogleAdsService.search_stream "
            "with a small customer query to attach names and currency.",
            "Reporting: Backend runs GAQL via GoogleAdsService.search_stream over user-selected date ranges. "
            "Campaign-level metrics are always fetched; additional slices (search terms, device, geography, ad group) "
            "are invoked when the analysis agent selects those tools for the user’s goal.",
            "Brief + synthesis: Raw aggregates are turned into a typed brief (deterministic structure). An LLM may add "
            "narrative synthesis; no LLM output is written back to Google Ads.",
        ],
    )
    doc.add_paragraph("User interface", style="Heading 3")
    add_bullets(
        doc,
        [
            "Reporting is interactive in the browser (scrollable report, optional synthesis sections).",
            "PDF export is not required for API compliance; current MVP focuses on on-screen report and local save.",
        ],
    )
    doc.add_paragraph("Data flow (conceptual diagram — text)", style="Heading 3")
    doc.add_paragraph(
        "Browser (Next.js UI) → HTTPS + API key → FastAPI /api → Google Ads API (read-only).\n"
        "Browser → OAuth 2.0 redirect → Google.\n"
        "FastAPI → Brief builder → optional goal-driven tools → Google Ads API."
    )

    doc.add_heading("5. API services called (read-only)", level=1)
    doc.add_paragraph(
        "All access is read via GoogleAdsService.search_stream (GAQL) and CustomerService.list_accessible_customers. "
        "We do not call mutate services (no CampaignService.mutate, AdGroupAdService, KeywordPlanIdeaService, etc.)."
    )
    add_table(
        doc,
        ("Purpose", "Google Ads API usage"),
        [
            (
                "List accounts user can access",
                "CustomerService.list_accessible_customers; then GoogleAdsService.search_stream on customer "
                "(id, name, currency, time zone, manager flag).",
            ),
            (
                "Campaign performance",
                "GoogleAdsService.search_stream — GAQL FROM campaign with segments.date, metrics (clicks, impressions, "
                "cost, conversions, conversion value), campaign.advertising_channel_type. Excludes REMOVED campaigns.",
            ),
            (
                "Search terms (top by spend, capped)",
                "GoogleAdsService.search_stream — FROM search_term_view.",
            ),
            (
                "Device segmentation",
                "GoogleAdsService.search_stream — FROM campaign with segments.device.",
            ),
            ("Geography", "GoogleAdsService.search_stream — FROM geographic_view."),
            ("Ad group rollups", "GoogleAdsService.search_stream — FROM ad_group."),
        ],
    )
    doc.add_paragraph()
    doc.add_paragraph(
        "Campaign types: We do not filter by channel in code; any non-removed campaign returned by reporting is in scope "
        "(Search, Performance Max, Display, Shopping, Video, etc., per Google’s classification)."
    )

    doc.add_heading("6. Security and compliance (summary)", level=1)
    add_bullets(
        doc,
        [
            "Transport: HTTPS between client and backend.",
            "Authentication to our API: X-API-Key on /api/... routes (except health and OAuth redirect endpoints as configured).",
            "Secrets: Developer token, OAuth client secret, and LLM keys are environment variables on the server, not committed to source control.",
            "Logging: We avoid logging full OAuth refresh tokens; operational logs may record customer id and errors for support.",
            "Token handling: Align deployed behavior with Google’s OAuth and Ads API policies; prefer server-stored refresh tokens per end-user account for production.",
        ],
    )

    doc.add_heading("7. Tool mockups / screenshots", level=1)
    doc.add_paragraph(
        "Required for externally accessible tools: embed 3–6 screenshots or mock-ups. Below is a prototype of the "
        "primary reporting UI; add live captures of the connection and generate flows when you want a fuller set."
    )
    doc.add_heading("7.1 Prototype — Paid Ads Performance Report (demo)", level=2)
    doc.add_paragraph(
        "This mock-up shows how users view Google Ads performance in the product: date window, headline ROAS with "
        "week-over-week context, 7-day spend sparkline, ROAS by campaign bars (Search vs Display), CAC / Spend / "
        "Conversions cards with target and period comparisons, and a signals area for notable issues."
    )
    if IMAGE_PATH.is_file():
        doc.add_picture(str(IMAGE_PATH), width=Inches(6.2))
    else:  # pragma: no cover
        doc.add_paragraph(f"[Image not found: {IMAGE_PATH}]")
    doc.add_paragraph(
        "Figure: Prototype report surface. Production UI is data-driven from API-backed briefs; layout and metrics match this experience.",
        style="Caption",
    )

    doc.add_heading("7.2 Additional captures (optional)", level=2)
    add_table(
        doc,
        ("#", "Screen", "Route / area"),
        [
            ("1", "Connections — Google Ads connect / status", "/connections"),
            ("2", "Generate — data sources step", "/generate step 1"),
            ("3", "Generate — account + goal + date range", "/generate step 2"),
            ("4", "Generate — review before run", "/generate step 3"),
            ("5", "In-progress / generating state", "/generate during API call"),
            ("6", "Live report (optional second shot)", "/generate after success or /reports/[slug]"),
        ],
    )
    doc.add_paragraph(
        "When building a PDF for Google, paste or embed additional PNGs so reviewers do not depend on Markdown rendering."
    )

    doc.add_heading("8. Declaration alignment (checklist)", level=1)
    add_bullets(
        doc,
        [
            "Capabilities: Reporting (read-only). Not: account/campaign creation or management via API; not Keyword Planner; "
            "not App Conversion Tracking / Remarketing API unless you add them.",
            "Token use with someone else’s tool: No (token is for Duct’s own backend).",
            "Campaign types: List major types or state “all types returned in campaign reporting.”",
        ],
    )

    doc.add_paragraph()
    p = doc.add_paragraph("End of document.")
    p.runs[0].italic = True

    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(DOCX_PATH))
    print(f"Wrote {DOCX_PATH}")

    if platform.system() == "Darwin":
        tr = subprocess.run(
            [
                "textutil",
                "-convert",
                "doc",
                str(DOCX_PATH),
                "-output",
                str(DOC_PATH),
            ],
            capture_output=True,
            text=True,
            check=False,
        )
        if tr.returncode == 0 and DOC_PATH.is_file():
            print(f"Wrote {DOC_PATH} (via macOS textutil)")
        elif tr.stderr:
            print(tr.stderr, file=sys.stderr)

    if not DOC_PATH.is_file():
        soffice = subprocess.run(
            ["which", "soffice"],
            capture_output=True,
            text=True,
            check=False,
        )
        if soffice.returncode == 0 and soffice.stdout.strip():
            bin_path = soffice.stdout.strip()
            outdir = DOCX_PATH.parent
            rv = subprocess.run(
                [
                    bin_path,
                    "--headless",
                    "--convert-to",
                    "doc",
                    "--outdir",
                    str(outdir),
                    str(DOCX_PATH),
                ],
                capture_output=True,
                text=True,
                check=False,
            )
            if rv.returncode == 0 and DOC_PATH.is_file():
                print(f"Wrote {DOC_PATH} (via LibreOffice)")
            else:
                print(
                    "LibreOffice conversion skipped or failed. Use Microsoft Word: File → Save As → "
                    "Word 97-2003 (.doc) if the uploader requires .doc.",
                    file=sys.stderr,
                )
        else:
            print(
                "No legacy .doc yet: on macOS, textutil should have run; otherwise install LibreOffice or open "
                "the .docx in Microsoft Word → Save As → Word 97-2003 Document (*.doc).",
                file=sys.stderr,
            )


if __name__ == "__main__":
    main()
